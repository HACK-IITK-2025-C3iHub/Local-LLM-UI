"""DOCX generation module for creating formatted Word documents."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


def create_docx_report(content, output_path, title="Policy Analysis Report"):
    """Generate formatted DOCX from text content.
    
    Args:
        content: Text content to convert to DOCX
        output_path: Path to save the DOCX file
        title: Document title
    """
    
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = title
    doc.core_properties.author = "Local LLM Policy Analyzer"
    doc.core_properties.subject = "NIST CSF Policy Analysis"
    
    # Define custom styles
    styles = doc.styles
    
    # Title style
    if 'CustomTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(26, 26, 26)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    if 'CustomHeading1' not in [s.name for s in styles]:
        h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.size = Pt(14)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(44, 62, 80)
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(10)
    
    # Heading 2 style
    if 'CustomHeading2' not in [s.name for s in styles]:
        h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.size = Pt(12)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(52, 73, 94)
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(8)
    
    # Body style
    if 'CustomBody' not in [s.name for s in styles]:
        body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Process content line by line
    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # Skip separator lines
        if re.match(r'^[=\-]{3,}$', line):
            continue
        
        # Major titles (all uppercase, long lines)
        if line.isupper() and len(line) > 10 and not line.startswith((' ', '\t', '-', '*', '•')):
            p = doc.add_paragraph(line, style='CustomTitle')
            continue
        
        # Section headers (ends with colon, not too long)
        if line.endswith(':') and len(line) < 80 and not line.startswith((' ', '\t')):
            p = doc.add_paragraph(line, style='CustomHeading2')
            continue
        
        # Bold text (**text**)
        if '**' in line:
            p = doc.add_paragraph(style='CustomHeading2')
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                m = re.match(r'^\*\*(.+)\*\*$', part)
                if m:
                    run = p.add_run(m.group(1))
                    run.bold = True
                else:
                    p.add_run(part)
            continue
        
        # Bullet points
        if line.strip().startswith(('-', '*', '•')):
            clean_line = line.strip().lstrip('-*• ')
            p = doc.add_paragraph(clean_line, style='List Bullet')
            continue
        
        # Numbered lists
        if re.match(r'^\s*\d+\.', line):
            clean_line = re.sub(r'^\s*\d+\.\s*', '', line)
            p = doc.add_paragraph(clean_line, style='List Number')
            continue
        
        # Regular paragraphs
        if line.strip():
            p = doc.add_paragraph(line, style='CustomBody')
        else:
            # Empty line - add spacing
            doc.add_paragraph()
    
    # Save document
    doc.save(output_path)


def generate_all_docx(results, output_base):
    """Generate DOCX versions of all analysis reports.
    
    Args:
        results: Dictionary containing all analysis results
        output_base: Base path for output files (without extension)
    
    Returns:
        List of generated DOCX file paths
    """
    
    from datetime import datetime
    
    docx_files = []
    
    # Gap Analysis DOCX
    gap_docx = f"{output_base}_gap_analysis.docx"
    create_docx_report(results['gap_analysis'], gap_docx, "Gap Analysis Report")
    docx_files.append(gap_docx)
    
    # Revised Policy DOCX
    policy_docx = f"{output_base}_revised_policy.docx"
    create_docx_report(results['revised_policy'], policy_docx, "Revised Policy Document")
    docx_files.append(policy_docx)
    
    # Roadmap DOCX
    roadmap_docx = f"{output_base}_roadmap.docx"
    create_docx_report(results['roadmap'], roadmap_docx, "Implementation Roadmap")
    docx_files.append(roadmap_docx)
    
    # Executive Summary DOCX
    exec_docx = f"{output_base}_executive_summary.docx"
    create_docx_report(results['executive_summary'], exec_docx, "Executive Summary")
    docx_files.append(exec_docx)
    
    # Comprehensive Report DOCX
    vuln_section = ''
    if 'vulnerability_analysis' in results:
        vuln_section = f"""
{'='*80}
SECURITY VULNERABILITY ANALYSIS
{'='*80}

{results['vulnerability_analysis']}
"""
    
    comprehensive = f"""
{'='*80}
COMPREHENSIVE POLICY ANALYSIS REPORT
{'='*80}
Policy: {results['policy_name']}
Analysis Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
Framework: NIST Cybersecurity Framework (CIS MS-ISAC 2024)
{'='*80}

{results['executive_summary']}

{'='*80}
DETAILED GAP ANALYSIS
{'='*80}

{results['gap_analysis']}{vuln_section}

{'='*80}
REVISED POLICY DOCUMENT
{'='*80}

{results['revised_policy']}

{'='*80}
IMPLEMENTATION ROADMAP
{'='*80}

{results['roadmap']}

{'='*80}
END OF REPORT
{'='*80}
"""
    
    comprehensive_docx = f"{output_base}_comprehensive_report.docx"
    create_docx_report(comprehensive, comprehensive_docx, "Comprehensive Policy Analysis")
    docx_files.append(comprehensive_docx)
    
    return docx_files
